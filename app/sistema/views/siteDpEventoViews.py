from django.shortcuts import render, redirect
from sistema.models.dpEvento import DpEvento
from sistema.models.atividadeCategoria import AtividadeCategoria
from sistema.models.membroExecucao import MembroExecucao
from sistema.models.escola import Escola
from sistema.models.departamento import Departamento
from sistema.models.ensino import Ensino
from sistema.services.alfrescoApi import AlfrescoAPI
from sistema.models.atividade import Atividade
from sistema.models.galeria import Galeria
from sistema.models.imagem import Imagem
from sistema.models.ticket import Ticket
from sistema.models.anexo import Anexo
from itertools import chain
from sistema.models.departamento import Departamento
from sistema.models.alocacao import Alocacao
from django.contrib.auth.decorators import login_required
from PIL import Image
import requests
from docx.enum.text import WD_UNDERLINE, WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, Inches, RGBColor
import json
import os
from django.http import FileResponse, HttpResponse
import docx
from django.http import JsonResponse
from rest_framework.authtoken.models import Token
from sistema.serializers.dpEventoSerializer import DpEventoSerializer
from sistema.serializers.ensinoSerializer import EnsinoSerializer
from sistema.serializers.escolaSerializer import EscolaSerializer
from sistema.serializers.escolaSerializer import EscolaSerializer
from django.db.models import Prefetch
from collections import defaultdict
from docx.image.exceptions import UnrecognizedImageError
from docx import Document
from datetime import date
from collections import Counter
import xlsxwriter
from io import BytesIO
from xlsxwriter.workbook import Worksheet
from xlsxwriter.format import Format

@login_required(login_url="/auth-user/login-user")
def gerencia_dp_eventos(request):
    page_title = "Projetos"
    dp_eventos = DpEvento.objects.all()
    count = dp_eventos.count()
    return render(
        request,
        "dpEventos/gerenciar_dp_eventos.html",
        {"dp_eventos": dp_eventos, "count": count, "page_title": page_title},
    )


@login_required(login_url="/auth-user/login-user")
def dpEventoTable(request):
    tipo = request.GET.get("tipo")
    data_inicio = request.GET.get("data_inicio")
    data_fim = request.GET.get("data_fim")
    order_by = request.GET.get("order_by")
    token, created = Token.objects.get_or_create(user=request.user)
    headers = {"Authorization": "Token " + token.key}

    dpEventoResponse = requests.get(
        "http://localhost:8000/dp-eventos",
        params={
            "tipo": tipo,
            "data_inicio": data_inicio,
            "data_fim": data_fim,
            "order_by": order_by,
        },
        headers=headers,
    )
    dpEventoResponseStatusCode = dpEventoResponse.status_code
    dpEventoResponse = json.loads(dpEventoResponse.content.decode())

    dpEventos = dpEventoResponse
    return render(request, "dpEventos/dp_eventos_tabela.html", {"dpEventos": dpEventos})


@login_required(login_url="/auth-user/login-user")
def dpEventoModal(request):
    id = request.GET.get("id")
    dpEvento = None
    data = {}
    escolas = Escola.objects.all()
    ensinos = Ensino.objects.all()
    data["escolas"] = EscolaSerializer(escolas, many=True).data
    data["ensinos"] = EnsinoSerializer(ensinos, many=True).data
    data["ct_emprestimo"] = DpEvento.EMPRESTIMO
    data["selectedEscolas"] = []
    if id:
        dpEvento = DpEvento.objects.get(id=id)
        data["dpEvento"] = DpEventoSerializer(dpEvento).data
        if dpEvento.acaoEnsino:
            ensinoSelected = dpEvento.acaoEnsino.id
            data["selected_ensino"] = (
                ensinoSelected if type(ensinoSelected) == "int" else int(ensinoSelected)
            )
        data["selectedEscolas"] = EscolaSerializer(
            dpEvento.escolas.all(), many=True
        ).data
    return render(request, "dpEventos/dp_eventos_modal.html", data)


@login_required(login_url="/auth-user/login-user")
def eliminarDpEvento(request, codigo):
    token, created = Token.objects.get_or_create(user=request.user)
    headers = {"Authorization": "Token " + token.key}

    dpEventoResponse = requests.delete(
        "http://localhost:8000/dp-eventos/" + str(codigo), headers=headers
    )
    dpEventoResponseStatusCode = dpEventoResponse.status_code
    dpEventoResponse = json.loads(dpEventoResponse.content.decode())

    return redirect("/gerencia_dp_eventos")


@login_required(login_url="/auth-user/login-user")
def saveDpEvento(request):
    token, created = Token.objects.get_or_create(user=request.user)
    headers = {"Authorization": "Token " + token.key}
    dpEventoData = json.loads(request.body)["data"]
    dpEventoResponse = requests.post(
        "http://localhost:8000/dp-eventos",
        json={"dpEvento": dpEventoData},
        headers=headers,
    )
    dpEventoResponseStatusCode = dpEventoResponse.status_code
    dpEventoResponse = json.loads(dpEventoResponse.content.decode())
    dpEvento = DpEvento.objects.get(id=dpEventoResponse["id"])

    return JsonResponse(dpEventoResponse, status=dpEventoResponseStatusCode)


@login_required(login_url="/auth-user/login-user")
def editarDpEvento(request, codigo):
    token, created = Token.objects.get_or_create(user=request.user)
    headers = {"Authorization": "Token " + token.key}
    body = json.loads(request.body)["data"]
    response = requests.put(
        "http://localhost:8000/dp-eventos/" + str(codigo), json=body, headers=headers
    )
    return JsonResponse(json.loads(response.content), status=response.status_code)


@login_required(login_url="/auth-user/login-user")
def dp_eventosSelect(request):
    dp_eventos = DpEvento.objects.all()
    return render(
        request, "dpEventos/dp_evento_select.html", {"dp_eventos": dp_eventos}
    )


@login_required(login_url="/auth-user/login-user")
def visualizarDpEvento(request, codigo):
    dpEvento = DpEvento.objects.prefetch_related(
        Prefetch(
            "membroexecucao_set",
            queryset=MembroExecucao.objects.select_related("itinerario")
            .prefetch_related("ticket_set")
            .prefetch_related("itinerario__itinerarioitem_set"),
        ),
    )
    dpEvento = dpEvento.get(id=codigo)

    departamentos = Departamento.objects.all()
    path_back = "gerencia_dp_eventos"
    dpEvento = DpEventoSerializer(dpEvento).data
    categoriaAtividades = AtividadeCategoria.objects.all()
    return render(
        request,
        "dpEventos/visualizar_dp_evento.html",
        {
            "dpEvento": dpEvento,
            "path_back": path_back,
            "departamentos": departamentos,
            "categoriaAtividades": categoriaAtividades,
        },
    )


def getFilteredEventos(filters, formatType="type 1"):
    eventos = None
    if "departamento_id" in filters and filters["departamento_id"]:
        eventos = DpEvento.objects.prefetch_related(
            Prefetch(
                "atividade_set",
                queryset=Atividade.objects.select_related("tipoAtividade").filter(
                    departamento=filters["departamento_id"], atividade_meta=True
                ),
            )
        )
    else:
        eventos = DpEvento.objects.prefetch_related(
            Prefetch(
                "atividade_set",
                queryset=Atividade.objects.select_related("tipoAtividade"),
            )
        ).filter(tipo=DpEvento.CURSO_GPS)

    if "data_inicio" in filters:
        eventos = eventos.filter(data_inicio__gte=filters["data_inicio"])
    if "data_fim" in filters:
        eventos = eventos.filter(data_fim__lte=filters["data_fim"])
    if "tipo" in filters and filters["tipo"]:
        eventos = eventos.filter(tipo=filters["tipo"])
    result = {}
    eventos = eventos.order_by("data_inicio")
    if formatType == "type 2":
        return eventos
    for evento in eventos:
        if evento.atividade_set.count() == 0:
            continue
        tipo = evento.tipo_formatado
        result.setdefault(tipo, {})

        for atividade in evento.atividade_set.all():
            tipo_atividade = atividade.tipoAtividade.nome if atividade.tipoAtividade else atividade.nome
            result[tipo].setdefault(tipo_atividade, [])
            result[tipo][tipo_atividade].append(atividade)

    report = defaultdict(list)
    for dp_evento in eventos:
        report[dp_evento.tipo_formatado].append(dp_evento)
    if True:
        return report
    return result


def getSectionTitle(doc, nomeEvento):
    title = doc.add_paragraph()
    title_run = title.add_run(f"{nomeEvento}")
    title_run.bold = True
    title_run.underline = WD_UNDERLINE.SINGLE
    title.space_after = Pt(0)
    return doc


def getCidade(doc, atividade):
    cidade = atividade.cidade
    cidadeParagraph = doc.add_paragraph()
    cidadeParagraph.add_run(f"cidade:").bold = True
    cidadeTexto = cidade.nome if cidade else "Não Informado"
    cidadeParagraph.add_run(f" {cidadeTexto}")
    cidadeParagraphFormat = cidadeParagraph.paragraph_format
    cidadeParagraphFormat.space_after = Pt(0)
    return doc


def getMatriculas(acaoEnsino):
    alocacoes = acaoEnsino.alocacao_set.all()
    matriculas = 0
    for alocacao in alocacoes:
        matriculas += (
            alocacao.quantidade_matriculas if alocacao.quantidade_matriculas else 0
        )
    return matriculas


def getServicosAtendimentos(atividade):
    servicos = atividade.servico_set.all()
    atendimentos = 0
    for servico in servicos:
        atendimentos += (
            servico.quantidadeAtendimentos if servico.quantidadeAtendimentos else 0
        )
    return atendimentos


def getQuantitativo(doc, atividade):
    quantitativoValor = atividade.tipo_quantitativo_valor
    quantitativoLabel = atividade.tipo_quantitativo_label

    if atividade.evento.acaoEnsino:
        quantitativoValor = getMatriculas(atividade.evento.acaoEnsino)
        quantitativoLabel = "Quantidade de Matrículas"
    if atividade.servico_set.count() > 0:
        quantitativoValor = getServicosAtendimentos(atividade)

    quantitativoParagraph = doc.add_paragraph()
    quantitativoParagraph.add_run(f"{quantitativoLabel}:").bold = True
    quantitativoParagraph.add_run(f" {quantitativoValor}")
    quantitativoParagraphFormat = quantitativoParagraph.paragraph_format
    quantitativoParagraphFormat.space_after = Pt(0)
    return doc


def getData(doc, atividade):
    dataInicio = atividade.data_realizacao_inicio_formatada
    dataFim = atividade.data_realizacao_fim_formatada
    dataStr = (
        f"{dataInicio} até {dataFim}" if dataInicio != dataFim else f"{dataInicio}"
    )
    dataParagraph = doc.add_paragraph()
    dataParagraph.add_run(f"Data:").bold = True
    dataParagraph.add_run(f" {dataStr}")
    dataParagraphFormat = dataParagraph.paragraph_format
    dataParagraphFormat.space_after = Pt(0)
    return doc


def getLocal(doc, atividade):
    local = atividade.endereco_completo
    localParagraph = doc.add_paragraph()
    localParagraph.add_run(f"Local:").bold = True
    localParagraph.add_run(f" {local}")
    localParagraphFormat = localParagraph.paragraph_format
    localParagraphFormat.space_after = Pt(0)
    return doc


def getEtapa(doc, atividade):
    etapa = atividade.evento.acaoEnsino.etapa_formatada
    if len(etapa) > 0:
        etapaParagraph = doc.add_paragraph()
        etapaParagraph.add_run(f"Etapa:").bold = True
        etapaParagraph.add_run(f" {etapa}")
        etapaParagraphFormat = etapaParagraph.paragraph_format
        etapaParagraphFormat.space_after = Pt(0)
        return doc

    return doc


def getSubAtividades(doc: Document, atividade):
    eventoEnsino = DpEvento.objects.filter(id=atividade.evento.id).first()
    if eventoEnsino.acaoEnsino:
        alocacoes = Alocacao.objects.filter(acaoEnsino=eventoEnsino.acaoEnsino)
        if alocacoes:
            subAtividadesParagraph = doc.add_paragraph()
            subAtividadesParagraph.add_run(f"Cursos Ofertados:").bold = True
            subAtividadesParagraphFormat = subAtividadesParagraph.paragraph_format
            subAtividadesParagraphFormat.space_after = Pt(0)
            for alocacao in alocacoes:
                curso = alocacao.curso.nome
                codigoSiga = alocacao.codigo_siga if alocacao.codigo_siga else ""
                quantidadeMatriculas = (
                    alocacao.quantidade_matriculas
                    if alocacao.quantidade_matriculas
                    else ""
                )
                subAtividadeParagraph = doc.add_paragraph(
                    f"{codigoSiga} {curso}: {quantidadeMatriculas}"
                )
                subAtividadeParagraphFormat = subAtividadeParagraph.paragraph_format
                subAtividadeParagraphFormat.space_after = Pt(0)
            return doc
    if atividade.servico_set.count() > 0:
        subAtividadesParagraph = doc.add_paragraph()
        subAtividadesParagraph.add_run(f"Serviços Ofertados:").bold = True
        subAtividadesParagraphFormat = subAtividadesParagraph.paragraph_format
        subAtividadesParagraphFormat.space_after = Pt(0)
        for servico in atividade.servico_set.all():
            servicoParagraph = doc.add_paragraph(
                f"{servico.nome}: {servico.quantidadeAtendimentos}"
            )
            servicoParagraphFormat = servicoParagraph.paragraph_format
            servicoParagraphFormat.space_after = Pt(0)
        return doc
    return doc


def getAtividadeLabel(doc, atividade, counter):
    cargaHoraria = atividade.cargaHoraria if atividade.cargaHoraria else ""
    atividadeLabel = doc.add_paragraph()
    atividadeLabel.add_run(f"Ação {counter} - {cargaHoraria}").bold = True
    atividadeLabelFormat = atividadeLabel.paragraph_format
    atividadeLabelFormat.space_after = Pt(0)
    return doc

def getAtividadeImage(doc: Document, atividade, counter):
    imagem = atividade.galeria.imagem_set.first()
    if not imagem:
        return doc
    
    alfresco_api = AlfrescoAPI()
    image_path = f"tmp/{imagem.id_alfresco}.jpg"
    alfresco_api.getNodeContent(image_path, imagem.id_alfresco)

    try:
        img = Image.open(image_path)
        img.save(image_path)
    except IOError:
        print("Error: Unable to read or save the image using Pillow")

    p = doc.add_paragraph()
    p.add_run(f"Figura {counter}: {imagem.descricao}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    try:
        img_paragraph = doc.add_paragraph()
        img_run = img_paragraph.add_run()
        img_run.add_picture(image_path, width=Inches(4.0))
        img_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except UnrecognizedImageError:
        error_message = (
            f"Error: Unable to recognize the image format for {imagem.descricao}."
        )
        p = doc.add_paragraph()
        p.add_run(error_message)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    finally:
        try:
            os.remove(image_path)
        except OSError as e:
            print(f"Error: Unable to delete the image file: {e}")

    return doc


def getAtividade(doc, atividade, counter):
    doc = getAtividadeLabel(doc, atividade, counter)
    doc = getCidade(doc, atividade)
    doc = getQuantitativo(doc, atividade)
    doc = getData(doc, atividade)
    doc = getLocal(doc, atividade)
    if atividade.evento.acaoEnsino:
        doc = getEtapa(doc, atividade)
    doc = getSubAtividades(doc, atividade)
    doc.add_paragraph()
    doc = getAtividadeImage(doc, atividade, counter)
    doc.add_paragraph()
    return doc


def reportEventos(eventos):
    eventosLength = len(eventos) > 0
    activitiesLength = 0
    for evento in eventos:
        activitiesLength = activitiesLength + evento.atividade_set.count()
    return activitiesLength > 0 and eventosLength > 0


def getRelatorioType1(doc, relatorioData):
    counter = 0
    for nomeEvento, eventos in relatorioData.items():
        if not reportEventos(eventos):
            continue

        doc = getSectionTitle(doc, nomeEvento)
        old_evento = 0
        for evento in eventos:
            current_evento = evento.id
            for atividade in evento.atividade_set.all():
                if current_evento != old_evento:
                    tipoTexto = evento.tipo_formatado if evento.tipo else "Não Informado"
                    cidadeNomeTexto = evento.cidade.nome if evento.cidade else "Não Informado"
                    doc = getSectionTitle(doc, f"{tipoTexto} - {cidadeNomeTexto}")
                    old_evento = current_evento
                tipoAtividadeTexto = atividade.tipoAtividade.nome if atividade.tipoAtividade else "Não Informado"
                doc = getSectionTitle(doc, f"{tipoAtividadeTexto}")
                counter = counter + 1
                doc = getAtividade(doc, atividade, counter)
    return doc


def getRelatorioType2(doc, relatorioData):
    # Initialize a Counter for status
    status_counter = Counter()
    tipo_counter = Counter()
    for evento in relatorioData:
        # calculate status based on data_inicio and data_fim
        status = "Não Foi possível determinar o status do evento"
        if evento.data_inicio and evento.data_fim:
            today = date.today()
            if today < evento.data_inicio:
                status = "Não Iniciado"
            elif evento.data_inicio <= today <= evento.data_fim:
                status = "Em Andamento"
            else:
                status = "Concluído"

        tipo_counter[evento.tipo_formatado] += 1

        # Update the status counter
        status_counter[status] += 1
    # total number of DpEventos
    doc.add_heading(f"Total de Eventos: {len(relatorioData)}", level=1)
    # Add status count to the top of the document
    doc.add_heading("Contagem por status do evento:", level=1)
    for status, count in status_counter.items():
        doc.add_paragraph(f"{status}: {count}")

    doc.add_paragraph("\n")
    doc.add_heading("Contagem por tipo do evento:", level=1)
    for tipo, count in tipo_counter.items():
        doc.add_paragraph(f"{tipo}: {count}")
    doc.add_paragraph("\n")
    for evento in relatorioData:
        status = "Não Foi possível determinar o status do evento"
        if evento.data_inicio and evento.data_fim:
            today = date.today()
            if today < evento.data_inicio:
                status = "Não Iniciado"
            elif evento.data_inicio <= today <= evento.data_fim:
                status = "Em Andamento"
            else:
                status = "Concluído"

        # Add each DpEvento's details to the document
        doc.add_paragraph(f"Tipo: {evento.tipo_formatado}")
        doc.add_paragraph(f"Status: {status}")
        doc.add_paragraph(f"Data de Início: {evento.data_inicio_formatada}")
        doc.add_paragraph(f"Data de Fim: {evento.data_fim_formatada}")
        doc.add_paragraph(
            f'Cidade: {evento.cidade.nome if evento.cidade else "N/A"}'
        )  # assuming 'nome' is the field name for city's name in Cidade model
        doc.add_paragraph(
            f'Escola: {evento.escola.nome if evento.escola else "N/A"}'
        )  # assuming 'nome' is the field name for school's name in Escola model
        doc.add_paragraph(f"Endereço: {evento.logradouro}")

        # Add a line break between each DpEvento
        doc.add_paragraph("\n")

    return doc


def createRelatorio(doc, relatorioData, type="type 1"):
    if type == "type 1":
        return getRelatorioType1(doc, relatorioData)
    if type == "type 2":
        return getRelatorioType2(doc, relatorioData)
    return doc


@login_required(login_url="/auth-user/login-user")
def relatorioDpEvento(request):
    filters = {}
    departamentoNome = ""
    if request.GET.get("departamento_id"):
        filters["departamento_id"] = request.GET.get("departamento_id")
        departamento = Departamento.objects.get(id=filters["departamento_id"])
        departamentoNome = departamento.nome
    if request.GET.get("data_inicio"):
        filters["data_inicio"] = request.GET.get("data_inicio")
    if request.GET.get("data_fim"):
        filters["data_fim"] = request.GET.get("data_fim")
    if request.GET.get("tipo"):
        filters["tipo"] = request.GET.get("tipo")

    relatorioTipo = "type 2" if departamentoNome == "Eventos" else "type 1"
    doc = docx.Document()
    relatorioData = getFilteredEventos(filters, relatorioTipo)
    doc = createRelatorio(doc, relatorioData, relatorioTipo)
    with open("tmp/dp_evento_report.docx", "wb") as f:
        doc.save(f)

    response = FileResponse(open("tmp/dp_evento_report.docx", "rb"))
    response[
        "Content-Type"
    ] = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    response["Content-Disposition"] = 'attachment; filename="dp_evento_report.docx"'

    os.remove("tmp/dp_evento_report.docx")
    return response


def getEventoRow(
    worksheet: Worksheet, evento: DpEvento, startRow: int, endRow: int, column: int, style: Format
) -> Worksheet:
    if startRow == endRow:
        worksheet.write(startRow, column, str(evento.tipo_formatado), style)
        return worksheet
    
    worksheet.merge_range(startRow, column, endRow, column, str(evento.tipo_formatado), style)
    return worksheet

def getAtividadeCountRow(
    worksheet: Worksheet, atividadeCount: int, startRow: int, endRow: int, column: int, style: Format
) -> Worksheet:
    if startRow == endRow:
        worksheet.write(startRow, column, atividadeCount, style)
        return worksheet

    worksheet.merge_range(startRow, column, endRow, column, atividadeCount, style)
    return worksheet

def getTipoAtividadeRow(
    worksheet: Worksheet, atividade: Atividade, startRow: int, endRow: int, column: int, style: Format
) -> Worksheet:
    if atividade.tipoAtividade:
        if startRow == endRow:
            worksheet.write(startRow, column, str(atividade.tipoAtividade.nome), style)
            return worksheet
        worksheet.merge_range(startRow, column, endRow, column, str(atividade.tipoAtividade.nome), style)
        return worksheet
    else:
        if startRow == endRow:
            worksheet.write(startRow, column, "N/I", style)
        worksheet.merge_range(startRow, column, endRow, column, "N/I", style)
    return worksheet

def getAtividadeHorasRow(
    worksheet: Worksheet, atividade: Atividade, startRow: int, endRow: int, column: int, style: Format
) -> Worksheet:
    if atividade.cargaHoraria:
        if startRow == endRow:
            worksheet.write(startRow, column, f"{str(int(atividade.cargaHoraria))}h", style)
            return worksheet
        worksheet.merge_range(startRow, column, endRow, column, f"{str(int(atividade.cargaHoraria))}h", style)
    else:
        if startRow == endRow:
            worksheet.write(startRow, column, "N/I", style)
        worksheet.merge_range(startRow, column, endRow, column, "N/I", style)
    return worksheet

def getAtividadeDescricaoRow(
    worksheet: Worksheet, atividade: Atividade, startRow: int, endRow: int, column: int, style: Format
) -> Worksheet:
    if atividade.descricao:
        if startRow == endRow:
            worksheet.write(startRow, column, str(atividade.descricao), style)
            return worksheet
        worksheet.merge_range(startRow, column, endRow, column, str(atividade.descricao), style)
    else:
        if startRow == endRow:
            worksheet.write(startRow, column, "N/I", style)
        worksheet.merge_range(startRow, column, endRow, column, "N/I", style)
    return worksheet

def getAtividadeEscolaRow(
    worksheet: Worksheet, atividade: Atividade, startRow: int, endRow: int, column: int, style: Format
) -> Worksheet:
    escola = atividade.evento.escolas.first()
    if escola:
        if startRow == endRow:
            worksheet.write(startRow, column, str(escola.nome), style)
            return worksheet
        worksheet.merge_range(startRow, column, endRow, column, str(escola.nome), style)
    else:
        if startRow == endRow:
            worksheet.write(startRow, column, "N/I", style)
        worksheet.merge_range(startRow, column, endRow, column, "N/I", style)
    return worksheet

def getAtividadeEventoEtapa(
    worksheet: Worksheet, atividade: Atividade, startRow: int, endRow: int, column: int, style: Format
) -> Worksheet:
    if atividade.evento.acaoEnsino:
        if atividade.evento.acaoEnsino.etapa:
            if startRow == endRow:
                worksheet.write(startRow, column, f"Etapa {atividade.evento.acaoEnsino.etapa}", style)
                return worksheet
            worksheet.merge_range(startRow, column, endRow, column, f"Etapa {atividade.evento.acaoEnsino.etapa}", style)
        else:
            if startRow == endRow:
                worksheet.write(startRow, column, "N/I", style)
            worksheet.merge_range(startRow, column, endRow, column, "N/I", style)
    else:
        if startRow == endRow:
            worksheet.write(startRow, column, "N/I", style)
        worksheet.merge_range(startRow, column, endRow, column, "N/I", style)
    return worksheet

def getAtividadeLocal(
    worksheet: Worksheet, atividade: Atividade, startRow: int, endRow: int, column: int, style: Format
) -> Worksheet:
    if len(atividade.evento.endereco_completo) > 0:
        if startRow == endRow:
            worksheet.write(startRow, column, str(atividade.evento.endereco_completo), style)
            return worksheet
        worksheet.merge_range(startRow, column, endRow, column, str(atividade.evento.endereco_completo), style)
    else:
        if startRow == endRow:
            worksheet.write(startRow, column, "N/I", style)
        worksheet.merge_range(startRow, column, endRow, column, "N/I", style)
    return worksheet

def getAtividadeDataInicio(
    worksheet: Worksheet, atividade: Atividade, startRow: int, endRow: int, column: int, style: Format
) -> Worksheet:
    if atividade.data_realizacao_inicio_formatada:
        if startRow == endRow:
            worksheet.write(startRow, column, atividade.data_realizacao_inicio_formatada, style)
            return worksheet
        worksheet.merge_range(startRow, column, endRow, column, atividade.data_realizacao_inicio_formatada, style)
    else:
        if startRow == endRow:
            worksheet.write(startRow, column, "N/I", style)
        worksheet.merge_range(startRow, column, endRow, column, "N/I", style)
    return worksheet

def getAtividadeDataFim(
    worksheet: Worksheet, atividade: Atividade, startRow: int, endRow: int, column: int, style: Format
) -> Worksheet:
    if atividade.data_realizacao_fim_formatada:
        if startRow == endRow:
            worksheet.write(startRow, column, atividade.data_realizacao_fim_formatada, style)
            return worksheet
        worksheet.merge_range(startRow, column, endRow, column, atividade.data_realizacao_fim_formatada, style)
    else:
        if startRow == endRow:
            worksheet.write(startRow, column, "N/I", style)
        worksheet.merge_range(startRow, column, endRow, column, "N/I", style)

    return worksheet

def getAtividadeAtendimentosRow(
    worksheet: Worksheet, atividade: Atividade, startRow: int, endRow: int, column: int, style: Format, alocacoes: list
) -> Worksheet:
    quantidade_atendimentos = atividade.quantidadeAtendimentos
    if not quantidade_atendimentos:
        quantidade_atendimentos = 0
        if len(alocacoes) > 0:
            for alocacao in alocacoes:
                if alocacao.quantidade_matriculas:
                    quantidade_atendimentos += alocacao.quantidade_matriculas
    if not quantidade_atendimentos or quantidade_atendimentos == 0:
        quantidade_atendimentos = atividade.quantidadeCertificacoes
    if not quantidade_atendimentos or quantidade_atendimentos == 0:
        quantidade_atendimentos = atividade.quantidadeInscricoes
    if not quantidade_atendimentos or quantidade_atendimentos == 0:
        quantidade_atendimentos = atividade.quantidadeMatriculas

    if quantidade_atendimentos:
        if startRow == endRow:
            worksheet.write(startRow, column, quantidade_atendimentos, style)
            return worksheet
        worksheet.merge_range(startRow, column, endRow, column, quantidade_atendimentos, style)
    else:
        if startRow == endRow:
            worksheet.write(startRow, column, "N/I", style)
        worksheet.merge_range(startRow, column, endRow, column, "N/I", style)

    return worksheet

def getAtividadeCurso(
    worksheet: Worksheet, alocacao: Alocacao, row: int, column: int, style: Format
) -> Worksheet:
    if alocacao.curso.nome:
        worksheet.write(row, column, alocacao.curso.nome, style)
    else:
        worksheet.write(row, column, "N/I", style)
    return worksheet

def getAtividadeCursoSiga(
    worksheet: Worksheet, alocacao: Alocacao, row: int, column: int, style: Format
) -> Worksheet:
    if alocacao.codigo_siga:
        worksheet.write(row, column, alocacao.codigo_siga, style)
    else:
        worksheet.write(row, column, "N/I", style)
    return worksheet

def getAtividadeCursoOficio(
    worksheet: Worksheet, alocacao: Alocacao, row: int, column: int, style: Format
) -> Worksheet:
    numero_oficio = alocacao.acaoEnsino.numero_oficio
    if numero_oficio:
        worksheet.write(row, column, numero_oficio, style)
    else:
        worksheet.write(row, column, "N/I", style)
    return worksheet

def getAtividadeCursoAtendimentos(
    worksheet: Worksheet, alocacao: Alocacao, row: int, column: int, style: Format
) -> Worksheet:
    if alocacao.quantidade_matriculas:
        worksheet.write(row, column, alocacao.quantidade_matriculas, style)
    else:
        worksheet.write(row, column, "N/I", style)
    return worksheet

def getAtividadeCursoTurno(
    worksheet: Worksheet, alocacao: Alocacao, row: int, column: int, style: Format
) -> Worksheet:
    turno = alocacao.turnos.first()
    if turno:
        worksheet.write(row, column, turno.nome, style)
    else:
        worksheet.write(row, column, "N/I", style)
    return worksheet

@login_required(login_url="/auth-user/login-user")
def relatorioSintetico(request):
    #  {'data_inicio': ['2023-05-26'], 'data_fim': ['2023-05-31'], 'departamento_id': ['1'], 'tipo': ['emprestimo']}
    data_inicio = request.GET.get("data_inicio")
    data_fim = request.GET.get("data_fim")
    departamento_id = request.GET.get("departamento_id")
    tipo = request.GET.get("tipo")

    output = BytesIO()
    workbook = xlsxwriter.Workbook(output)
    worksheet = workbook.add_worksheet()
    merge_format = workbook.add_format(
        {
            "bold": 1,
            "border": 1,
            "align": "center",
            "valign": "vcenter",
            "fg_color": "yellow",
        }
    )

    centered = workbook.add_format()
    centered.set_align('center')
    centered.set_align('vcenter')
    centered.set_text_wrap(True)

    headers = [
        'Evento',
        "Quantidade de Ações",
        "Ação de Extensão",
        "Horas Executadas",
        "Quantidade Atendimentos",
        "Descrição da Ação",
        "COTEC/RESPONSÁVEL",
        "Etapa",
        "Local",
        "Data de Início (Atividade)",
        "Data de Fim (Atividade)",
        "Curso",
        "COD. SIGA",
        "Turno",
        "Ofício",
        "ALUNOS"
    ]

        
    max_len_headers = [len(header) for header in headers]

    for i, header in enumerate(headers):
        worksheet.write(0, i, header, centered)
    filters = {}
    eventos = DpEvento.objects
    if data_inicio:
        filters["data_inicio"] = data_inicio
    if data_fim:
        filters["data_fim"] = data_fim
    if departamento_id:
        filters["departamento_id"] = departamento_id
    if tipo:
        filters["tipo"] = tipo

    eventos = getFilteredEventos(filters, "type 2")
    
    startRow = 1
    previousRowCount = 0
    atividadeCounter = 0
    for evento in eventos:
        atividades = Atividade.objects.filter(evento=evento)
        
        alocacoes = []
        if evento.acaoEnsino:
            acaoEnsino = evento.acaoEnsino
            alocacoes = Alocacao.objects.filter(acaoEnsino=acaoEnsino).select_related('curso').all()

        alocacoesCount = 1 if len(alocacoes) == 0 else len(alocacoes)
        atividadeCount = len(atividades)
        if atividadeCount == 0:
            continue

        row = startRow

        # coluna do evento
        startEventoRow = row
        endEventoRow = (alocacoesCount * atividadeCount) + (previousRowCount)
        worksheet = getEventoRow(worksheet, evento, startEventoRow, endEventoRow, headers.index("Evento"), centered)

        # coluna da atividade
        atividadeStartRow = startRow
        atividadeEndRow = startRow + alocacoesCount - 1
        for atividade in atividades:
            atividadeCounter += 1
            worksheet = getAtividadeCountRow(worksheet, atividadeCounter, atividadeStartRow, atividadeEndRow, headers.index("Quantidade de Ações"), centered)
            worksheet = getTipoAtividadeRow(worksheet, atividade, atividadeStartRow, atividadeEndRow, headers.index("Ação de Extensão"), centered)
            worksheet = getAtividadeHorasRow(worksheet, atividade, atividadeStartRow, atividadeEndRow, headers.index("Horas Executadas"), centered)
            worksheet = getAtividadeAtendimentosRow(worksheet, atividade, atividadeStartRow, atividadeEndRow, headers.index("Quantidade Atendimentos"), centered, alocacoes)
            worksheet = getAtividadeDescricaoRow(worksheet, atividade, atividadeStartRow, atividadeEndRow, headers.index("Descrição da Ação"), centered)
            worksheet = getAtividadeEscolaRow(worksheet, atividade, atividadeStartRow, atividadeEndRow, headers.index("COTEC/RESPONSÁVEL"), centered)
            worksheet = getAtividadeEventoEtapa(worksheet, atividade, atividadeStartRow, atividadeEndRow, headers.index("Etapa"), centered)
            worksheet = getAtividadeLocal(worksheet, atividade, atividadeStartRow, atividadeEndRow, headers.index("Local"), centered)
            worksheet = getAtividadeDataInicio(worksheet, atividade, atividadeStartRow, atividadeEndRow, headers.index("Data de Início (Atividade)"), centered)
            worksheet = getAtividadeDataFim(worksheet, atividade, atividadeStartRow, atividadeEndRow, headers.index("Data de Fim (Atividade)"), centered)
            atividadeStartRow += alocacoesCount
            atividadeEndRow += alocacoesCount

            # coluna do curso
            if len(alocacoes) == 0:
                worksheet.write(row, headers.index("Curso"), "N/A", centered)
                worksheet.write(row, headers.index("COD. SIGA"), "N/A", centered)
                worksheet.write(row, headers.index("Turno"), "N/A", centered)
                worksheet.write(row, headers.index("Ofício"), "N/A", centered)
                worksheet.write(row, headers.index("ALUNOS"), "N/A", centered)
                row += 1

            for alocacao in alocacoes:
                worksheet = getAtividadeCurso(worksheet, alocacao, row, headers.index("Curso"), centered)
                worksheet = getAtividadeCursoSiga(worksheet, alocacao, row, headers.index("COD. SIGA"), centered)
                worksheet = getAtividadeCursoTurno(worksheet, alocacao, row, headers.index("Turno"), centered)
                worksheet = getAtividadeCursoOficio(worksheet, alocacao, row, headers.index("Ofício"), centered)
                worksheet = getAtividadeCursoAtendimentos(worksheet, alocacao, row, headers.index("ALUNOS"), centered)
                row += 1

        startRow = row
        previousRowCount = row - 1

    for i, header in enumerate(headers):
        if header == "Descrição da Ação" or header == "Local":
            worksheet.set_column(i, i, 50)
        elif header == "Curso" or header == "COD. SIGA" or header == "COTEC/RESPONSÁVEL":
            worksheet.set_column(i, i, 30)
        else:
            worksheet.set_column(i, i, max_len_headers[i] + 5)

    workbook.close()

    response = HttpResponse(output.getvalue(), content_type="application/vnd.ms-excel")
    response["Content-Disposition"] = "attachment; filename=RelatorioSintetico.xlsx"
    return response



def addHeading(document, text, style):
    heading = document.add_paragraph()
    run = heading.add_run(text)
    run.font.size = Pt(style.get("size"))
    run.font.underline = style.get("underline")
    heading.alignment = style.get("alignment")
    return document

def addParagraph(document, text, style):
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.font.size = Pt(style.get("size"))
    paragraph.alignment = style.get("alignment")
    return document

def addTable(document, rowsCount, colsCount, rowsContent, title):
    table = document.add_table(rows=rowsCount+1, cols=colsCount)
    table.style = "Table Grid"
    if type(title) is str:
        title_cell = table.rows[0].cells[0].merge(table.rows[0].cells[1])
        title_cell.text = title
        title_cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    if type(title) is list:
        for i, t in enumerate(title):
            title_cell = table.rows[0].cells[i]
            title_cell.text = t
            title_cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for i in range(0, rowsCount):
        for j in range(colsCount):
            table.rows[i+1].cells[j].text = rowsContent[i][j]

    return document

def processImagemRelatorioEvento(doc: Document, imagem: Imagem, anexo: Anexo):
    id_alfresco = imagem.id_alfresco if imagem else anexo.id_alfresco
    descricao = imagem.descricao if imagem else ""

    if not id_alfresco:
        return doc

    alfresco_api = AlfrescoAPI()
    image_path = f"tmp/{id_alfresco}.jpg"
    alfresco_api.getNodeContent(image_path, id_alfresco)

    try:
        img = Image.open(image_path)
        img.save(image_path)
    except IOError:
        print("Error: Unable to read or save the image using Pillow")

    try:
        img_paragraph = doc.add_paragraph()
        img_run = img_paragraph.add_run()
        img_run.add_picture(image_path, width=Inches(4.0))
        img_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_paragraph.paragraph_format.space_after = Pt(0)
    except UnrecognizedImageError:
        error_message = (
            f"Error: Unable to recognize the image format for {descricao}."
        )
        p = doc.add_paragraph()
        p.add_run(error_message)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    finally:
        try:
            os.remove(image_path)
        except OSError as e:
            print(f"Error: Unable to delete the image file: {e}")

    return doc

@login_required(login_url="/auth-user/login-user")
def relatorioPorEvento(request, evento_id):
    evento = DpEvento.objects.get(id=evento_id) 
    doc = docx.Document()
    headingStyle = {
        "size": 16,
        "alignment": WD_ALIGN_PARAGRAPH.CENTER,
        "underline": False,
    }

    paragraphStyle = {
        "size": 12,
        "alignment": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }

    figuraCounter = 1
    tabelaCounter = 1
    heading_text = f"prestação de contas - {evento.tipo_formatado} {evento.edicao}º edição".upper()
    doc = addHeading(doc, heading_text, headingStyle)
    doc = addParagraph(doc, f"{evento.descricao}", paragraphStyle)
    local = ""
    
    escolas = evento.escolas.all()
    for escola in escolas:
        local += f"{escola.nome}\n"

    tableData = [
        ["Nome do Evento", f"{evento.tipo_formatado}"],
        ["Data", f"de {evento.data_inicio.strftime('%d/%m/%Y')} até {evento.data_fim.strftime('%d/%m/%Y')}"],
        ["Horário", f"{evento.horarioInicio.strftime('%H:%M')} até {evento.horarioFim.strftime('%H:%M')}"],
        ["Local", local],
    ]
    
    rowsCount = len(tableData)
    colsCount = len(tableData[0])
    doc = addTable(doc, rowsCount, colsCount, tableData, "FICHA TÉCNICA")
    captionText = f"Tabela {tabelaCounter}: Ficha técnica do {evento.tipo_formatado}"
    if evento.edicao:
        captionText += f" {evento.edicao}º edição"
    caption = doc.add_paragraph(captionText, 'Caption')
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(0)
    doc.add_paragraph()
    for run in caption.runs:
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    tabelaCounter += 1

    headingStyle["alignment"] = WD_ALIGN_PARAGRAPH.LEFT
    departamentoComunicacao = Departamento.objects.get(nome__icontains="Comunicação")
    atividadesComunicacao = Atividade.objects.filter(evento=evento, departamento=departamentoComunicacao)
    
    doc.add_paragraph()
    doc = addHeading(doc, "ARTES DO EVENTO", headingStyle)
    doc.add_paragraph()
    
    for i, atividadeComunicacao in enumerate(atividadesComunicacao):
        anexos = Anexo.objects.filter(model="Atividade", id_model=atividadeComunicacao.id)
        for anexo in anexos:
            mimeTypeIsImage = "image" in anexo.mime_type or "jpeg" in anexo.mime_type or "png" in anexo.mime_type or "jpg" in anexo.mime_type
            if (mimeTypeIsImage):
                captionText = f"Figura {figuraCounter}: "
                if anexo.descricao:
                    captionText = f" {anexo.descricao}"
                else:
                    captionText += f"Nenhuma descrição disponível"
                doc = processImagemRelatorioEvento(doc, None, anexo)
                caption = doc.add_paragraph(captionText, 'Caption')
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption.paragraph_format.space_before = Pt(0)
                doc.add_paragraph()
                for run in caption.runs:
                    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
                figuraCounter += 1
            else:
                print("Anexo não é imagem", anexo.mime_type)

    doc = addHeading(doc, "IMAGENS DO EVENTO", headingStyle)
    doc.add_paragraph()
    geleriaGeral = Galeria.objects.filter(evento=evento, nome="galeria geral do evento").first()
    imagensGaleriaGeral = Imagem.objects.filter(galeria=geleriaGeral)
    for imagem in imagensGaleriaGeral:
        doc = processImagemRelatorioEvento(doc, imagem, None)
        captionText = f"Figura {figuraCounter}: "
        if imagem.descricao:
            captionText = f"{imagem.descricao}"
        else:
            captionText += f"Nenhuma descrição disponível."
        caption = doc.add_paragraph(captionText, 'Caption')
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.paragraph_format.space_before = Pt(0)
        doc.add_paragraph()
        for run in caption.runs:
            run.font.color.rgb = RGBColor(0x00, 0x00, 0x00) 
        figuraCounter += 1


    atividadeCategoriaProgramacao = AtividadeCategoria.objects.get(name__icontains="programação")
    atividadesProgramacao = Atividade.objects.filter(evento=evento, atividadeCategorias__id=atividadeCategoriaProgramacao.id)
    
    doc = addHeading(doc, "PROGRAMAÇÃO DO EVENTO", headingStyle)
    doc.add_paragraph()
    for i, atividadeProgramacao in enumerate(atividadesProgramacao):
        galeria = atividadeProgramacao.galeria
        imagens = Imagem.objects.filter(galeria=galeria)
        for imagem in imagens:
            captionText = f"Figura {figuraCounter}: "
            if imagem.descricao:
                captionText = f"{imagem.descricao}"
            else: 
                captionText += f"Nenhuma descrição disponível"
            doc = processImagemRelatorioEvento(doc, imagem, None)
            caption = doc.add_paragraph(captionText, 'Caption')
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption.paragraph_format.space_before = Pt(0)
            doc.add_paragraph()
            for run in caption.runs:
                run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
            figuraCounter += 1

    doc = addHeading(doc, "EXECUÇÂO FINANCEIRA", headingStyle)
    doc.add_paragraph()
    membro_execucao_objs = MembroExecucao.objects.filter(evento__id=evento.id)
    ensino_objs = Ensino.objects.filter(dpevento__id=evento.id)
    alocacao_objs = Alocacao.objects.filter(acaoEnsino__in=ensino_objs)
    atividade_objs = Atividade.objects.filter(evento__id=evento.id)

    tickets_membro_execucao = Ticket.objects.filter(membro_execucao__in=membro_execucao_objs)
    tickets_alocacao = Ticket.objects.filter(alocacao__in=alocacao_objs)
    tickets_atividade = Ticket.objects.filter(atividade__in=atividade_objs)

    demandas = list(chain(tickets_membro_execucao, tickets_alocacao, tickets_atividade))
    execucaoFinanceiraHeader = ["ID","Data da Demanda","solicitante", "Assunto", "Beneficiário/Destinatário", "Tipo", "Valor"]
    execucaoFinanceiraData = []
    for demanda in demandas:
        solicitante = demanda.solicitante.nome if demanda.solicitante else "Não informado"
        beneficiario = demanda.beneficiario.nome if demanda.beneficiario else None
        if not beneficiario:
            beneficiario = demanda.escola.nome if demanda.escola else None
        if not beneficiario:
            beneficiario = "Não informado"
        execucaoFinanceiraData.append([
            demanda.id_protocolo, 
            demanda.data_criacao_formatada, 
            solicitante,
            demanda.observacao, 
            beneficiario,
            demanda.tipo_formatado,
            str(demanda.valor_executado) if demanda.valor_executado else "Não informado"
        ])
    
    doc = addTable(doc, len(execucaoFinanceiraData), len(execucaoFinanceiraHeader), execucaoFinanceiraData, execucaoFinanceiraHeader)
    captionText = f"Tabela {tabelaCounter}: Execução financeira {evento.tipo_formatado}"
    if evento.edicao:
        captionText += f" {evento.edicao}º edição"
    caption = doc.add_paragraph(captionText, 'Caption')
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(0)
    doc.add_paragraph()
    for run in caption.runs:
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    tabelaCounter += 1

    valorTotalHeader = ["Escola", "Valor Total (R$)"]
    valorTotalData = {}
    for demanda in demandas:
        if demanda.escola.nome not in valorTotalData:
            valorTotalData[demanda.escola.nome] = 0
        
        valorTotalData[demanda.escola.nome] += demanda.valor_executado if demanda.valor_executado else 0
    
    doc.add_paragraph()
    dataValorTotal = []
    doc = addHeading(doc, "VALOR TOTAL POR ESCOLA", headingStyle)
    doc.add_paragraph()
    for key, value in valorTotalData.items():
        dataValorTotal.append([key, str(value)])
    
    for escola in evento.escolas.all():
        if escola.nome not in valorTotalData:
            dataValorTotal.append([escola.nome, "0"])
    
    doc = addTable(doc, len(dataValorTotal), len(valorTotalHeader), dataValorTotal, valorTotalHeader)
    captionText = f"Tabela {tabelaCounter}: Execução financeira - valor por escola {evento.tipo_formatado}"
    if evento.edicao:
        captionText += f" {evento.edicao}º edição"
    caption = doc.add_paragraph(captionText, 'Caption')
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(0)
    doc.add_paragraph()
    for run in caption.runs:
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    tabelaCounter += 1

    with open("tmp/relatorio_evento.docx", "wb") as f:
        doc.save(f)

    response = FileResponse(open("tmp/relatorio_evento.docx", "rb"))
    response[
        "Content-Type"
    ] = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    response["Content-Disposition"] = 'attachment; filename="relatorio_evento.docx"'

    os.remove("tmp/relatorio_evento.docx")
    return response